"""File parser module for extracting text content from various file formats."""
import base64
import json
from pathlib import Path
from typing import Optional
from loguru import logger

# Supported file extensions and their MIME types
SUPPORTED_EXTENSIONS = {
    # Plain text
    'txt': 'text/plain',
    'md': 'text/markdown',
    'markdown': 'text/markdown',
    'json': 'application/json',
    'csv': 'text/csv',
    'yaml': 'application/x-yaml',
    'yml': 'application/x-yaml',

    # Code files
    'py': 'text/x-python',
    'js': 'text/javascript',
    'ts': 'text/typescript',
    'tsx': 'text/typescript',
    'jsx': 'text/jsx',
    'java': 'text/x-java',
    'cpp': 'text/x-c++src',
    'c': 'text/x-csrc',
    'h': 'text/x-csrc',
    'go': 'text/x-go',
    'rs': 'text/x-rust',
    'rb': 'text/x-ruby',
    'php': 'text/x-php',
    'sql': 'text/x-sql',
    'sh': 'text/x-shellscript',
    'bash': 'text/x-shellscript',
    'dockerfile': 'text/x-dockerfile',
    'xml': 'text/xml',
    'html': 'text/html',
    'css': 'text/css',

    # Document formats
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def detect_file_type(filename: str) -> Optional[str]:
    """Detect file MIME type from extension."""
    ext = Path(filename).suffix.lstrip('.').lower()
    return SUPPORTED_EXTENSIONS.get(ext)


def is_supported_file(filename: str) -> bool:
    """Check if file type is supported."""
    return detect_file_type(filename) is not None


def parse_plain_text(content: bytes, filename: str) -> str:
    """Parse plain text files with automatic encoding detection."""
    try:
        import chardet
        detected = chardet.detect(content)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        text = content.decode(encoding, errors='replace')
        return text.strip()
    except Exception as e:
        logger.error(f"Error parsing plain text file {filename}: {e}")
        raise ValueError(f"Failed to parse text file: {str(e)}")


def parse_code_file(content: bytes, filename: str) -> str:
    """Parse code files (assume UTF-8)."""
    try:
        text = content.decode('utf-8', errors='replace')
        return text.strip()
    except Exception as e:
        logger.error(f"Error parsing code file {filename}: {e}")
        raise ValueError(f"Failed to parse code file: {str(e)}")


def parse_docx(content: bytes, filename: str) -> str:
    """Parse DOCX files using python-docx."""
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(content))
        paragraphs = []

        # Extract all paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_cells))
            if table_rows:
                paragraphs.extend(table_rows)

        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"Error parsing DOCX file {filename}: {e}")
        raise ValueError(f"Failed to parse Word document: {str(e)}")


def parse_xlsx(content: bytes, filename: str) -> str:
    """Parse XLSX files using openpyxl."""
    try:
        import openpyxl
        from io import BytesIO

        workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
        sheets_content = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_lines = [f"Sheet: {sheet_name}"]

            # Extract all rows
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else '' for cell in row]
                sheet_lines.append(" | ".join(row_values))

            sheets_content.append("\n".join(sheet_lines))

        return "\n\n".join(sheets_content).strip()
    except Exception as e:
        logger.error(f"Error parsing XLSX file {filename}: {e}")
        raise ValueError(f"Failed to parse Excel file: {str(e)}")


def parse_file(filename: str, mime_type: str, content_base64: str) -> str:
    """
    Main dispatcher for file parsing.

    Args:
        filename: Original filename
        mime_type: MIME type of the file
        content_base64: Base64-encoded file content

    Returns:
        Extracted text content

    Raises:
        ValueError: If file parsing fails or format is unsupported
    """
    try:
        # Decode base64 content
        content = base64.b64decode(content_base64)

        # Validate file size
        if len(content) > MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds {MAX_FILE_SIZE / 1024 / 1024:.0f}MB limit")

        # Check if supported
        if not is_supported_file(filename):
            ext = Path(filename).suffix.lstrip('.')
            raise ValueError(f"File type '.{ext}' is not supported")

        # Get extension
        ext = Path(filename).suffix.lstrip('.').lower()

        # Parse based on file type
        if ext in ['txt', 'md', 'markdown', 'json', 'csv', 'yaml', 'yml']:
            return parse_plain_text(content, filename)
        elif ext in ['docx']:
            return parse_docx(content, filename)
        elif ext in ['xlsx']:
            return parse_xlsx(content, filename)
        else:
            # Treat all other supported types as code/text files
            return parse_code_file(content, filename)

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Unexpected error parsing file {filename}: {e}")
        raise ValueError(f"Failed to parse file: {str(e)}")
